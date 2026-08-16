"""Materialize and evaluate deterministic document fixtures."""
from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Any


def _pdf_bytes(text: str) -> bytes:
    """Build a small deterministic text PDF without an external renderer."""

    def escaped(value: str) -> str:
        return (
            value.replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)")
        )

    lines = [line for line in text.splitlines() if line.strip()]
    operations = ["BT", "/F1 12 Tf", "72 720 Td", "14 TL"]
    for index, line in enumerate(lines):
        if index:
            operations.append("T*")
        operations.append(f"({escaped(line)}) Tj")
    operations.append("ET")
    stream = "\n".join(operations).encode("latin-1", errors="replace")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> "
            b"/Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        (
            f"<< /Length {len(stream)} >>\nstream\n".encode("ascii")
            + stream
            + b"\nendstream"
        ),
    ]
    payload = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for index, body in enumerate(objects, 1):
        offsets.append(len(payload))
        payload.extend(f"{index} 0 obj\n".encode("ascii"))
        payload.extend(body)
        payload.extend(b"\nendobj\n")
    xref = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    payload.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        payload.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    payload.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(payload)


def _materialize_document_fixture(
    case: dict[str, Any],
    target: Path,
) -> None:
    mode = str(case.get("fixture_mode") or "")
    text = str(case.get("fixture_text") or "")
    target.parent.mkdir(parents=True, exist_ok=True)
    if target.suffix.casefold() == ".docx":
        if mode == "malformed-not-zip":
            target.write_bytes(b"not a zip package")
            return
        if mode == "missing-document-part":
            import zipfile

            with zipfile.ZipFile(target, "w") as archive:
                archive.writestr("[Content_Types].xml", "<Types/>")
            return
        from docx import Document

        document = Document()
        if "header" in mode:
            document.sections[0].header.paragraphs[0].text = (
                "Candidate | candidate@example.com"
            )
        document.add_heading("SUMMARY", level=1)
        paragraph = document.add_paragraph()
        first, separator, remainder = text.partition("Supported ")
        if first:
            paragraph.add_run(first)
        paragraph.add_run("Supported ").bold = "mixed-runs" in mode
        paragraph.add_run(remainder if separator else text).italic = (
            "styled" in mode
        )
        document.add_heading("EXPERIENCE", level=1)
        if "table" in mode:
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Role"
            table.cell(0, 1).text = "Evidence"
            table.cell(1, 0).text = "Analyst"
            table.cell(1, 1).text = "Completed 2 years of professional experience."
        else:
            document.add_paragraph(
                "Completed 2 years of professional experience.",
                style="List Bullet",
            )
        if "repeated-text" in mode:
            document.add_paragraph("Supported repeated validation text.")
            document.add_paragraph("Supported repeated validation text.")
        if "footer" in mode:
            document.sections[0].footer.paragraphs[0].text = "Fixture footer"
        document.save(target)
        return
    if target.suffix.casefold() == ".pdf":
        target.write_bytes(_pdf_bytes(text))
        return
    target.write_text(text, encoding="utf-8")


def _evaluate_document_case(
    case: dict[str, Any],
    *,
    root: Path,
) -> dict[str, Any]:
    from .ingestion import ExtractionError, load

    expected_parse = bool(case.get("expected_parse"))
    try:
        with tempfile.TemporaryDirectory() as directory:
            fixture_name = Path(str(case["fixture"])).name
            fixture = Path(directory) / fixture_name
            _materialize_document_fixture(case, fixture)
            loaded = load(fixture)
    except (ExtractionError, OSError, ValueError, KeyError) as exc:
        return {
            "id": str(case["id"]),
            "passed": not expected_parse,
            "parsed": False,
            "detail": f"{type(exc).__name__}: {exc}",
        }
    fragment = str(case.get("expected_text_fragment") or "")
    text = str(loaded.get("text") or "")
    return {
        "id": str(case["id"]),
        "passed": expected_parse and (not fragment or fragment in text),
        "parsed": True,
        "detail": "",
    }
