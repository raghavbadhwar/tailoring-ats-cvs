"""Structure-preserving DOCX patch operations.

The patcher edits only the exact approved span inside an existing paragraph and
saves a duplicate document. Paragraph style, run formatting, tables, headers,
footers, relationships, and unmodified package parts are retained by
``python-docx``. Complex fields or hyperlinks that cannot be represented by
ordinary runs are blocked rather than flattened silently.
"""
from __future__ import annotations

from collections.abc import Iterable, Iterator
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph


class DocxPatchError(ValueError):
    """Raised when an approved edit cannot be applied without ambiguity."""


def _iter_table_paragraphs(table: Table) -> Iterator[Paragraph]:
    seen_cells: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            cell_key = id(cell._tc)
            if cell_key in seen_cells:
                continue
            seen_cells.add(cell_key)
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_story_paragraphs(story: Any) -> Iterator[Paragraph]:
    yield from story.paragraphs
    for table in story.tables:
        yield from _iter_table_paragraphs(table)


def iter_document_paragraphs(document: DocumentType) -> Iterator[Paragraph]:
    """Yield body, table, header, and footer paragraphs without duplicates."""

    seen: set[int] = set()

    def unique(paragraphs: Iterable[Paragraph]) -> Iterator[Paragraph]:
        for paragraph in paragraphs:
            key = id(paragraph._p)
            if key in seen:
                continue
            seen.add(key)
            yield paragraph

    yield from unique(_iter_story_paragraphs(document))
    for section in document.sections:
        for story in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            yield from unique(_iter_story_paragraphs(story))


def _run_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def _replace_across_runs(
    paragraph: Paragraph,
    start: int,
    end: int,
    replacement: str,
) -> None:
    runs = paragraph.runs
    if not runs:
        raise DocxPatchError("matching DOCX paragraph has no editable runs")

    run_text = _run_text(paragraph)
    if run_text != paragraph.text:
        raise DocxPatchError(
            "matching DOCX paragraph contains complex fields or hyperlinks; "
            "use rebuild mode or simplify the paragraph before applying"
        )

    offsets: list[tuple[int, int]] = []
    cursor = 0
    for run in runs:
        next_cursor = cursor + len(run.text)
        offsets.append((cursor, next_cursor))
        cursor = next_cursor

    start_run = next(
        (index for index, (left, right) in enumerate(offsets) if left <= start < right),
        None,
    )
    end_position = end - 1
    end_run = next(
        (
            index
            for index, (left, right) in enumerate(offsets)
            if left <= end_position < right
        ),
        None,
    )
    if start_run is None or end_run is None:
        raise DocxPatchError("approved span could not be mapped to DOCX runs")

    start_left, _ = offsets[start_run]
    end_left, _ = offsets[end_run]
    prefix = runs[start_run].text[: start - start_left]
    suffix = runs[end_run].text[end - end_left :]

    if start_run == end_run:
        runs[start_run].text = prefix + replacement + suffix
        return

    runs[start_run].text = prefix + replacement
    for index in range(start_run + 1, end_run):
        runs[index].text = ""
    runs[end_run].text = suffix


def patch_docx(
    source: Path,
    output: Path,
    changes: Iterable[dict[str, Any]],
) -> None:
    """Apply exact, unique approved replacements to a duplicate DOCX file."""

    document = Document(str(source))
    paragraphs = list(iter_document_paragraphs(document))

    for change in changes:
        expected = str(change.get("expected_text", ""))
        replacement = str(change.get("replacement_text", ""))
        if not expected:
            raise DocxPatchError(f"change {change.get('id')} has no expected text")

        matches: list[tuple[Paragraph, int]] = []
        for paragraph in paragraphs:
            text = paragraph.text
            cursor = 0
            while True:
                index = text.find(expected, cursor)
                if index < 0:
                    break
                matches.append((paragraph, index))
                cursor = index + max(1, len(expected))

        if not matches:
            raise DocxPatchError(
                f"change {change.get('id')} expected text was not found in DOCX"
            )
        if len(matches) > 1:
            raise DocxPatchError(
                f"change {change.get('id')} is ambiguous in DOCX: {len(matches)} matches"
            )

        paragraph, start = matches[0]
        _replace_across_runs(paragraph, start, start + len(expected), replacement)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
