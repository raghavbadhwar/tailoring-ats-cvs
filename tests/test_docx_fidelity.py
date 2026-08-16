from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from ats_agent.documents import patch_document


class DocxFidelityTests(unittest.TestCase):
    def test_partial_replacement_preserves_unaffected_run_formatting(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "resume.docx"
            output = root / "tailored.docx"

            document = Document()
            paragraph = document.add_paragraph(style="List Bullet")
            prefix = paragraph.add_run("Prefix ")
            prefix.bold = True
            target = paragraph.add_run("Helped build")
            target.italic = True
            middle = paragraph.add_run(" automated order workflows")
            middle.underline = True
            suffix = paragraph.add_run(" with 42 tests.")
            suffix.bold = True
            document.save(source)

            patch_document(
                source,
                output,
                [
                    {
                        "id": "C1",
                        "operation": "replace_span",
                        "anchor": {
                            "part": "word/document.xml",
                            "paragraph_index": 0,
                        },
                        "expected_text": "Helped build",
                        "replacement_text": "Contributed to building",
                    }
                ],
                mode="preserve",
            )

            tailored = Document(output)
            result = tailored.paragraphs[0]
            self.assertEqual(
                result.text,
                "Prefix Contributed to building automated order workflows "
                "with 42 tests.",
            )
            self.assertEqual(result.style.name, "List Bullet")
            self.assertEqual(result.runs[0].text, "Prefix ")
            self.assertTrue(result.runs[0].bold)
            self.assertIn("Contributed to building", result.runs[1].text)
            self.assertTrue(result.runs[1].italic)
            self.assertEqual(
                result.runs[2].text,
                " automated order workflows",
            )
            self.assertTrue(result.runs[2].underline)
            self.assertEqual(result.runs[3].text, " with 42 tests.")
            self.assertTrue(result.runs[3].bold)

    def test_insertion_inside_table_cell_keeps_table_structure(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "resume.docx"
            output = root / "tailored.docx"

            document = Document()
            table = document.add_table(rows=1, cols=1)
            paragraph = table.cell(0, 0).paragraphs[0]
            paragraph.text = "PROJECTS"
            document.save(source)

            patch_document(
                source,
                output,
                [
                    {
                        "id": "C1",
                        "operation": "insert_after",
                        "anchor": {
                            "part": "word/document.xml",
                            "paragraph_index": 0,
                        },
                        "expected_text": "PROJECTS",
                        "replacement_text": (
                            "Built Python workflow automation with 42 tests."
                        ),
                    }
                ],
                mode="preserve",
            )

            tailored = Document(output)
            self.assertEqual(len(tailored.tables), 1)
            cell_paragraphs = tailored.tables[0].cell(0, 0).paragraphs
            self.assertEqual(cell_paragraphs[0].text, "PROJECTS")
            self.assertEqual(
                cell_paragraphs[1].text,
                "Built Python workflow automation with 42 tests.",
            )

    def test_non_body_part_edit_fails_explicitly(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "resume.docx"
            output = root / "tailored.docx"

            document = Document()
            document.sections[0].header.paragraphs[0].text = (
                "candidate@example.com"
            )
            document.add_paragraph("PROJECTS")
            document.save(source)

            with self.assertRaisesRegex(ValueError, "unsupported DOCX part"):
                patch_document(
                    source,
                    output,
                    [
                        {
                            "id": "C1",
                            "operation": "replace_span",
                            "anchor": {
                                "part": "word/header1.xml",
                                "paragraph_index": 0,
                            },
                            "expected_text": "candidate@example.com",
                            "replacement_text": "new@example.com",
                        }
                    ],
                    mode="preserve",
                )

            self.assertFalse(output.exists())


if __name__ == "__main__":
    unittest.main()
