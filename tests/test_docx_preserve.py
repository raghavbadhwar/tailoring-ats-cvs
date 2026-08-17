import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from ats_agent.documents import patch_document
from ats_agent.ingestion import extract
from ats_agent.workflow import apply_manifest, build_proposal


class DocxPreserveTests(unittest.TestCase):
    def test_apply_preserves_run_formatting_header_and_source(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "resume.docx"
            job = root / "job.md"
            proposal_path = root / "proposal.json"
            output = root / "tailored.docx"

            document = Document()
            document.sections[0].header.paragraphs[0].text = (
                "candidate@example.com"
            )
            document.add_heading("SUMMARY", level=1)
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(
                "Helped build automated order workflows with 42 tests."
            )
            run.bold = True
            run.font.name = "Aptos"
            document.save(source)
            original_bytes = source.read_bytes()

            job.write_text(
                "Workflow automation experience is required.",
                encoding="utf-8",
            )
            proposal = build_proposal(
                source,
                job,
                candidate_id="candidate-a",
            )
            supported = [
                change
                for change in proposal["changes"]
                if change["supported"]
            ]
            self.assertEqual(len(supported), 1)
            proposal_path.write_text(json.dumps(proposal), encoding="utf-8")

            manifest = root / "approved.json"
            manifest.write_text(
                json.dumps(
                    {
                        "schema_version": 2,
                        "proposal": str(proposal_path),
                        "proposal_digest": proposal["proposal_digest"],
                        "approved_change_ids": [supported[0]["id"]],
                        "output": str(output),
                        "document_mode": "preserve",
                    }
                ),
                encoding="utf-8",
            )
            result = apply_manifest(manifest, [supported[0]["id"]])

            self.assertEqual(source.read_bytes(), original_bytes)
            self.assertEqual(result["document_mode"], "preserve")
            self.assertEqual(result["format_lock"]["status"], "not_requested")
            self.assertIn("validated_output", result["coverage"])
            self.assertGreater(result["coverage"]["validated_output"]["covered_term_count"], 0)
            tailored = Document(output)
            tailored_paragraph = next(
                item
                for item in tailored.paragraphs
                if "workflow automation" in item.text
            )
            self.assertEqual(tailored_paragraph.style.name, "List Bullet")
            self.assertTrue(tailored_paragraph.runs[0].bold)
            self.assertEqual(tailored_paragraph.runs[0].font.name, "Aptos")
            self.assertEqual(
                tailored.sections[0].header.paragraphs[0].text,
                "candidate@example.com",
            )
            self.assertIn("candidate@example.com", extract(output))

    def test_pdf_output_is_rejected_without_renderer(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "resume.txt"
            job = root / "job.md"
            source.write_text(
                "PROJECTS\n"
                "- Helped build automated order workflows with 42 tests.\n",
                encoding="utf-8",
            )
            job.write_text(
                "Workflow automation experience is required.",
                encoding="utf-8",
            )
            proposal = build_proposal(source, job)
            supported = [
                change
                for change in proposal["changes"]
                if change["supported"]
            ]
            proposal_path = root / "proposal.json"
            proposal_path.write_text(json.dumps(proposal), encoding="utf-8")
            manifest = root / "approved.json"
            manifest.write_text(
                json.dumps(
                    {
                        "schema_version": 2,
                        "proposal": str(proposal_path),
                        "proposal_digest": proposal["proposal_digest"],
                        "approved_change_ids": [supported[0]["id"]],
                        "output": str(root / "fake.pdf"),
                    }
                ),
                encoding="utf-8",
            )
            with self.assertRaisesRegex(ValueError, "genuine PDF renderer"):
                apply_manifest(manifest, [supported[0]["id"]])

    def test_strict_preserve_rejects_paragraph_insertions(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source, output = root / "resume.docx", root / "out.docx"
            Document().save(source)
            with self.assertRaisesRegex(ValueError, "anchored text replacement only"):
                patch_document(
                    source,
                    output,
                    [{"id": "C1", "operation": "insert_after", "replacement_text": "New", "anchor": {}}],
                    mode="strict-preserve",
                )

    def test_strict_preserve_reports_a_verified_structural_lock(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source, job = root / "resume.docx", root / "job.md"
            proposal_path, output = root / "proposal.json", root / "out.docx"
            document = Document()
            document.add_paragraph("Helped build automated workflows with 42 tests.")
            document.save(source)
            job.write_text("Workflow automation is required.", encoding="utf-8")
            proposal = build_proposal(source, job)
            change = next(item for item in proposal["changes"] if item["supported"])
            proposal_path.write_text(json.dumps(proposal), encoding="utf-8")
            manifest = root / "approval.json"
            manifest.write_text(json.dumps({
                "schema_version": 2,
                "proposal": str(proposal_path),
                "proposal_digest": proposal["proposal_digest"],
                "approved_change_ids": [change["id"]],
                "output": str(output),
                "document_mode": "strict-preserve",
            }), encoding="utf-8")
            result = apply_manifest(manifest, [change["id"]])
            self.assertEqual(result["format_lock"]["status"], "verified")
            self.assertEqual(result["format_lock"]["rendered_layout"], "unverified")


if __name__ == "__main__":
    unittest.main()
